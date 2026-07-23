import logging
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table

logger = logging.getLogger(__name__)


def _extract_paragraph_text(document: DocumentObject) -> list[str]:
    """Extract non-empty paragraph text from a DOCX document."""

    return [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]


def _extract_table_text(table: Table) -> list[str]:
    """
    Extract table text.

    Resumes commonly use tables for layout, and table content is not included in
    document.paragraphs by python-docx.
    """

    rows: list[str] = []

    for row in table.rows:
        cells: list[str] = []

        for cell in row.cells:
            cell_text = " ".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )

            if cell_text:
                cells.append(cell_text)

        if cells:
            rows.append(" | ".join(cells))

    return rows


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract paragraphs and table content from a DOCX document.

    The service executes this synchronous parser in a worker thread.
    """

    if not file_path.is_file():
        raise FileNotFoundError(f"DOCX file does not exist: {file_path}")

    try:
        document = Document(file_path)
        parts = _extract_paragraph_text(document)

        for table in document.tables:
            parts.extend(_extract_table_text(table))
    except Exception as exc:
        logger.exception(
            "DOCX parsing failed",
            extra={"file_path": str(file_path)},
        )
        raise RuntimeError("Unable to parse the uploaded DOCX document") from exc

    extracted_text = "\n".join(parts).strip()

    if not extracted_text:
        raise ValueError("No readable text was found in the DOCX document")

    return extracted_text
